"""
=============================================================================
MODULE 8 - MODEL EVALUATION  |  report_generator.py
=============================================================================
Thesis-quality Word (.docx) report for held-out test set evaluation.
13 sections with embedded charts, data-driven narratives, clinical
interpretation, and deployment readiness assessment.
=============================================================================
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    MODULE_VERSION, REPORT, PRIMARY_METRIC, HIGH_PRIORITY_STATES,
    METRIC_DISPLAY, COMPARISON_METRICS, OVERFIT_THRESHOLD,
    CALIBRATION_GOOD_THRESHOLD, CLINICAL_COST_WEIGHTS,
)

log = logging.getLogger(__name__)


class M8ReportGenerator:
    """Build thesis-quality Word report from M8 evaluation results."""

    def __init__(self, verbose: bool = True):
        self.verbose = verbose
        self.doc = Document()
        self._setup_styles()

    def _log(self, msg: str):
        if self.verbose:
            print(f"  [M8-Report] {msg}")

    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        style.font.name = REPORT["font"]
        style.font.size = Pt(REPORT["size"])
        style.paragraph_format.space_after = Pt(4)
        for level in range(1, 4):
            hs = self.doc.styles[f"Heading {level}"]
            hs.font.name = REPORT["font"]
            hs.font.color.rgb = RGBColor(*REPORT["heading_rgb"])

    def _interp(self, text: str):
        """Add an interpretation paragraph."""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(REPORT["interp_pt"])
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)

    def _interp_bold_lead(self, bold_part: str, rest: str):
        """Add interpretation with bold lead-in."""
        p = self.doc.add_paragraph()
        rb = p.add_run(bold_part)
        rb.font.size = Pt(REPORT["interp_pt"])
        rb.bold = True
        rn = p.add_run(rest)
        rn.font.size = Pt(REPORT["interp_pt"])

    def _add_image(self, path: Path, width: float = None):
        """Embed image if it exists."""
        if path and path.exists():
            w = Inches(width or REPORT["img_md"])
            self.doc.add_picture(str(path), width=w)

    def _add_caption(self, text: str):
        """Add figure/table caption."""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(REPORT["caption_pt"])
        r.font.color.rgb = RGBColor(*REPORT["meta_rgb"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_table(self, df: pd.DataFrame, max_rows: int = 30):
        """Add a DataFrame as a Word table."""
        if df.empty:
            self._interp("No data available for this table.")
            return

        df_show = df.head(max_rows)
        table = self.doc.add_table(
            rows=len(df_show) + 1, cols=len(df_show.columns),
            style=REPORT["table_style"],
        )

        # Header
        for j, col in enumerate(df_show.columns):
            cell = table.rows[0].cells[j]
            cell.text = str(col)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(REPORT["table_pt"])
                    run.bold = True

        # Data
        for i, (_, row) in enumerate(df_show.iterrows()):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                if isinstance(val, float):
                    cell.text = f"{val:.4f}" if abs(val) < 10 else f"{val:.1f}"
                else:
                    cell.text = str(val) if val is not None else "--"
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(REPORT["table_pt"])

    # ══════════════════════════════════════════════════════════════════════
    #  Public API
    # ══════════════════════════════════════════════════════════════════════

    def build(
        self,
        output_dir: Path,
        eval_results: list,
        comparison_df: pd.DataFrame,
        gen_results: dict,
        gap_table: pd.DataFrame,
        gap_matrix: pd.DataFrame,
        cal_results: dict,
        clinical_results: dict,
        demo_results: dict,
        rob_results: dict,
        stat_comparison: Any,
        error_results: dict,
        prof_results: dict,
        test_data: Any,
        m7_meta: Any,
        split_dists: dict,
        plot_paths: Dict[str, Path],
    ) -> Path:
        """Build the full 13-section report."""
        self._log("Building report...")

        self._add_title_page(test_data, m7_meta)
        self._section_1_executive_summary(eval_results, gen_results, cal_results,
                                          clinical_results, stat_comparison)
        self._section_2_test_data(test_data, split_dists)
        self._section_3_performance(eval_results, comparison_df, plot_paths)
        self._section_4_generalization(gen_results, gap_table, gap_matrix, plot_paths)
        self._section_5_calibration(cal_results, plot_paths)
        self._section_6_clinical(clinical_results, plot_paths)
        self._section_7_demographics(demo_results, plot_paths)
        self._section_8_equity(demo_results, plot_paths)
        self._section_9_robustness(rob_results, plot_paths)
        self._section_10_statistical(stat_comparison, plot_paths)
        self._section_11_errors(error_results, plot_paths)
        self._section_12_inference(prof_results, plot_paths)
        self._section_13_summary(eval_results, gen_results, cal_results,
                                 clinical_results, rob_results, prof_results)

        report_path = output_dir / "M8_Model_Evaluation_Report.docx"
        self.doc.save(str(report_path))
        self._log(f"Report saved: {report_path.name}")
        return report_path

    # ══════════════════════════════════════════════════════════════════════
    #  Title Page
    # ══════════════════════════════════════════════════════════════════════

    def _add_title_page(self, test_data, m7_meta):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n\n\n")
        r = p.add_run("Model Evaluation Report")
        r.font.size = Pt(REPORT["title_pt"])
        r.font.color.rgb = RGBColor(*REPORT["heading_rgb"])
        r.bold = True

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Held-Out Test Set Analysis")
        r2.font.size = Pt(REPORT["subtitle_pt"])
        r2.font.color.rgb = RGBColor(*REPORT["subtitle_rgb"])

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = (
            f"Module 8 v{MODULE_VERSION}\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
            f"Test samples: {test_data.n_samples:,}\n"
            f"Classes: {test_data.n_classes}\n"
        )
        if m7_meta:
            meta_text += f"Best M7 model: {m7_meta.best_model}\n"
        r3 = p3.add_run(meta_text)
        r3.font.size = Pt(REPORT["size"])
        r3.font.color.rgb = RGBColor(*REPORT["meta_rgb"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 1 -- Executive Summary
    # ══════════════════════════════════════════════════════════════════════

    def _section_1_executive_summary(self, eval_results, gen_results,
                                     cal_results, clinical_results, stat_comp):
        self.doc.add_heading("1. Executive Summary", level=1)

        valid = [r for r in eval_results if not r.skipped]
        if not valid:
            self._interp("No models were successfully evaluated.")
            return

        best = valid[0]
        f1 = best.test_metrics.get("f1_weighted", 0)
        acc = best.test_metrics.get("accuracy", 0)

        self.doc.add_heading("1.1 Best Model on Held-Out Test Set", level=2)
        self._interp_bold_lead(
            f"{best.model_name} ",
            f"achieved the highest test-set performance with F1 (weighted) = {f1:.4f} "
            f"and accuracy = {acc:.4f}. This model was evaluated on data that was never "
            f"seen during training or hyperparameter tuning, providing an unbiased "
            f"estimate of real-world performance."
        )

        self.doc.add_heading("1.2 Key Findings", level=2)

        # Generalization
        if gen_results:
            overfit_count = sum(1 for g in gen_results.values()
                               if g.verdict in ("MILD_OVERFIT", "SEVERE_OVERFIT"))
            if overfit_count:
                self._interp_bold_lead(
                    "Overfitting detected: ",
                    f"{overfit_count}/{len(gen_results)} models show validation-to-test "
                    f"performance drops exceeding {OVERFIT_THRESHOLD:.0%}. "
                    f"This suggests these models may have partially memorised patterns "
                    f"in the training data rather than learning generalisable features."
                )
            else:
                self._interp_bold_lead(
                    "Good generalisation: ",
                    f"All {len(gen_results)} models maintained consistent performance "
                    f"from validation to test set, indicating robust learning."
                )

        # Calibration
        if cal_results:
            well_cal = sum(1 for c in cal_results.values()
                          if c.verdict == "WELL_CALIBRATED")
            self._interp_bold_lead(
                "Calibration: ",
                f"{well_cal}/{len(cal_results)} models are well-calibrated "
                f"(ECE < {CALIBRATION_GOOD_THRESHOLD:.0%}). Well-calibrated models "
                f"produce confidence scores that accurately reflect their true "
                f"probability of being correct -- critical for clinical decision support."
            )

        # Clinical
        if clinical_results:
            best_clin = min(clinical_results.values(), key=lambda c: c.cost_per_sample)
            self._interp_bold_lead(
                "Clinical utility: ",
                f"{best_clin.model_name} achieves the lowest clinical cost "
                f"({best_clin.cost_per_sample:.2f} per sample), reflecting the "
                f"asymmetric penalty for missing high-priority states (Fear, SIB, "
                f"Anger, Toilet) vs. false alarms."
            )

        self.doc.add_heading("1.3 Clinical Readiness Assessment", level=2)
        # Build a simple checklist
        checks = [
            ("Test-set F1 > 0.7", f1 > 0.7),
            ("No severe overfitting", not any(
                g.verdict == "SEVERE_OVERFIT" for g in gen_results.values()
            ) if gen_results else True),
            ("Calibrated (ECE < 5%)", any(
                c.verdict == "WELL_CALIBRATED" for c in cal_results.values()
            ) if cal_results else False),
        ]
        for label, passed in checks:
            symbol = "PASS" if passed else "FAIL"
            self._interp(f"  [{symbol}] {label}")

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 2 -- Test Data Overview
    # ══════════════════════════════════════════════════════════════════════

    def _section_2_test_data(self, test_data, split_dists):
        self.doc.add_heading("2. Test Data Overview", level=1)

        self._interp(
            f"The held-out test set contains {test_data.n_samples:,} samples across "
            f"{test_data.n_classes} classes. This data was separated from training and "
            f"validation sets at the data-splitting stage (Module 3) and has not been "
            f"used for any model development, hyperparameter tuning, or feature "
            f"selection decisions."
        )

        self.doc.add_heading("2.1 Class Distribution", level=2)
        if test_data.class_counts:
            counts = test_data.class_counts
            total = sum(counts.values())
            rows = []
            for cls in sorted(counts.keys()):
                c = counts[cls]
                rows.append({"Class": cls, "Count": c, "Proportion": f"{c/total:.1%}"})
            self._add_table(pd.DataFrame(rows))

        self.doc.add_heading("2.2 Train/Val/Test Distribution Comparison", level=2)
        if split_dists:
            self._interp(
                "Comparing class proportions across splits helps identify sampling "
                "drift that could bias evaluation. Stratified splitting should produce "
                "near-identical proportions across train, validation, and test."
            )
            rows = []
            for cls in sorted(set().union(*[d.keys() for d in split_dists.values()])):
                row = {"Class": cls}
                for split in ("train", "val", "test"):
                    dist = split_dists.get(split, {})
                    total = sum(dist.values()) if dist else 1
                    row[f"{split} (n)"] = dist.get(cls, 0)
                    row[f"{split} (%)"] = f"{dist.get(cls, 0)/total:.1%}" if total else "--"
                rows.append(row)
            self._add_table(pd.DataFrame(rows))

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 3 -- Test Set Performance
    # ══════════════════════════════════════════════════════════════════════

    def _section_3_performance(self, eval_results, comparison_df, plots):
        self.doc.add_heading("3. Test Set Performance", level=1)

        self._interp(
            "This section presents the core performance metrics computed on the "
            "held-out test set. Unlike validation metrics (which may have been used "
            "to select hyperparameters or early-stopping points), these test metrics "
            "provide an unbiased estimate of how each model would perform on truly "
            "unseen data from the same distribution."
        )

        self.doc.add_heading("3.1 Summary Comparison Table", level=2)
        self._add_table(comparison_df)
        self._add_image(plots.get("v04"), REPORT["img_lg"])

        valid = [r for r in eval_results if not r.skipped]
        if valid:
            best = valid[0]
            worst = valid[-1]
            self._interp_bold_lead(
                f"Overall ranking: ",
                f"{best.model_name} leads with F1={best.test_metrics.get('f1_weighted', 0):.4f}, "
                f"while {worst.model_name} trails at "
                f"F1={worst.test_metrics.get('f1_weighted', 0):.4f}. "
                f"The gap of {best.test_metrics.get('f1_weighted', 0) - worst.test_metrics.get('f1_weighted', 0):.4f} "
                f"suggests meaningful performance differences that are further tested "
                f"for statistical significance in Section 10."
            )

        self.doc.add_heading("3.2 Confusion Matrices", level=2)
        self._interp(
            "Confusion matrices reveal the pattern of predictions for each model. "
            "Diagonal elements represent correct classifications; off-diagonal "
            "elements are misclassifications. For clinical deployment, the "
            "off-diagonal pattern matters more than overall accuracy -- a model "
            "that confuses Fear with Anger may be acceptable, but one that "
            "confuses Fear with Baseline is clinically dangerous."
        )
        self._add_image(plots.get("v01"), REPORT["img_lg"])

        self.doc.add_heading("3.3 ROC Curves", level=2)
        self._interp(
            "Receiver Operating Characteristic (ROC) curves plot the true positive "
            "rate against the false positive rate at varying classification thresholds. "
            "The area under the curve (AUC) summarises discriminative ability: "
            "AUC = 1.0 indicates perfect discrimination, AUC = 0.5 indicates random "
            "classification. Each curve represents one class in a one-vs-rest scheme."
        )
        self._add_image(plots.get("v02"), REPORT["img_lg"])

        self.doc.add_heading("3.4 Precision-Recall Curves", level=2)
        self._interp(
            "Precision-recall curves are particularly informative for imbalanced "
            "datasets where some emotional states are rare. Average Precision (AP) "
            "summarises the curve. High AP for a rare class indicates the model can "
            "identify that state without flooding clinicians with false alarms."
        )
        self._add_image(plots.get("v03"), REPORT["img_lg"])

        self.doc.add_heading("3.5 Per-Class F1 Heatmap", level=2)
        self._interp(
            "This heatmap shows F1 scores for every model-class combination. "
            "Dark blue cells indicate strong performance; lighter cells indicate "
            "classes where the model struggles. Columns with uniformly light "
            "colours suggest inherently difficult-to-classify states."
        )
        self._add_image(plots.get("v05"), REPORT["img_md"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 4 -- Generalization Gap
    # ══════════════════════════════════════════════════════════════════════

    def _section_4_generalization(self, gen_results, gap_table, gap_matrix, plots):
        self.doc.add_heading("4. Generalisation Gap Analysis", level=1)

        self._interp(
            "The generalisation gap measures how much performance degrades from "
            "the validation set (used during model development) to the held-out "
            "test set (never seen). A large gap (>5%) suggests the model has "
            "overfit to patterns specific to the training/validation data. "
            "This is especially concerning in clinical settings where the model "
            "will encounter novel participants with different physiological baselines."
        )

        self.doc.add_heading("4.1 Val vs Test Metrics", level=2)
        self._add_table(gap_table)
        self._add_image(plots.get("v06"), REPORT["img_md"])

        self.doc.add_heading("4.2 Overfitting Heatmap", level=2)
        self._interp(
            "Positive values (red) indicate overfit -- the model performed better "
            "on validation than on the truly unseen test set. Negative values "
            "(green) indicate the model generalises better than expected. "
            f"Cells exceeding {OVERFIT_THRESHOLD:.0%} are flagged."
        )
        self._add_image(plots.get("v07"), REPORT["img_md"])

        self.doc.add_heading("4.3 Generalisation Waterfall", level=2)
        self._add_image(plots.get("v08"), REPORT["img_md"])

        # Verdicts
        if gen_results:
            for name, gr in gen_results.items():
                if gr.flags:
                    for flag in gr.flags:
                        self._interp_bold_lead(f"{name}: ", flag)

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 5 -- Calibration
    # ══════════════════════════════════════════════════════════════════════

    def _section_5_calibration(self, cal_results, plots):
        self.doc.add_heading("5. Confidence Calibration", level=1)

        self._interp(
            "A well-calibrated classifier produces probability estimates that "
            "accurately reflect the true likelihood of correctness. When a "
            "calibrated model says it is 80% confident in a prediction, it should "
            "be correct approximately 80% of the time. Calibration is essential "
            "for clinical deployment: a caregiver alert system must distinguish "
            "between high-confidence and uncertain predictions to avoid alarm "
            "fatigue while not missing genuine events."
        )

        self.doc.add_heading("5.1 Reliability Diagrams", level=2)
        self._interp(
            "Reliability diagrams plot predicted confidence (x-axis) against "
            "observed accuracy (y-axis). The diagonal represents perfect "
            "calibration. Points above the diagonal indicate underconfidence "
            "(the model is better than it thinks); points below indicate "
            "overconfidence (the model is worse than it claims)."
        )
        self._add_image(plots.get("v09"), REPORT["img_lg"])

        self.doc.add_heading("5.2 ECE & Brier Score Comparison", level=2)
        self._interp(
            "Expected Calibration Error (ECE) is the weighted average gap "
            "between confidence and accuracy across bins. Lower ECE means "
            "better calibration. The Brier score measures the mean squared "
            "difference between predicted probabilities and actual outcomes -- "
            "it penalises both miscalibration and poor discrimination."
        )
        self._add_image(plots.get("v10"), REPORT["img_md"])

        if cal_results:
            for name, cr in cal_results.items():
                if cr.ece is not None:
                    verdict_text = (
                        "well-calibrated" if cr.verdict == "WELL_CALIBRATED"
                        else "miscalibrated"
                    )
                    self._interp_bold_lead(
                        f"{name}: ",
                        f"ECE = {cr.ece:.4f}, Brier = {cr.brier_score:.4f} -- "
                        f"{verdict_text}. "
                        f"Mean confidence = {cr.mean_confidence:.3f}, "
                        f"overconfident in {cr.overconfidence_frac:.0%} of bins."
                    )

        self.doc.add_heading("5.3 Confidence Distribution", level=2)
        self._add_image(plots.get("v11"), REPORT["img_md"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 6 -- Clinical Utility
    # ══════════════════════════════════════════════════════════════════════

    def _section_6_clinical(self, clinical_results, plots):
        self.doc.add_heading("6. Clinical Utility Analysis", level=1)

        self._interp(
            "Standard ML metrics treat all errors equally, but in clinical "
            "practice, errors have asymmetric costs. Missing a Fear event in an "
            "autistic child (false negative) is far more dangerous than a false "
            "alarm that temporarily diverts a caregiver's attention. This section "
            "applies a clinical cost matrix that penalises false negatives for "
            "high-priority states (Fear, SIB, Anger, Toilet) more heavily than "
            "false positives."
        )

        cost_text = "Clinical cost weights: "
        for state, w in CLINICAL_COST_WEIGHTS.items():
            if state != "default":
                cost_text += f"{state} FN={w['fn_cost']}x, "
        cost_text += f"default FN={CLINICAL_COST_WEIGHTS['default']['fn_cost']}x."
        self._interp(cost_text)

        self.doc.add_heading("6.1 High-Priority False Negative Rates", level=2)
        self._interp(
            "The False Negative Rate (FNR) for high-priority states is the single "
            "most critical metric for clinical deployment. A high FNR means the "
            "model frequently misses dangerous states, leaving the child without "
            "timely intervention."
        )
        self._add_image(plots.get("v12"), REPORT["img_md"])

        self.doc.add_heading("6.2 Clinical Cost Heatmap", level=2)
        self._add_image(plots.get("v13"), REPORT["img_md"])

        self.doc.add_heading("6.3 Number Needed to Screen (NNS)", level=2)
        self._interp(
            "NNS estimates how many samples must be screened to detect one true "
            "event. Lower NNS indicates a more efficient screening tool. NNS = "
            "1/(Sensitivity x Prevalence). For rare but dangerous states like "
            "SIB, even a moderate NNS may be acceptable given the severity."
        )
        self._add_image(plots.get("v14"), REPORT["img_md"])

        self.doc.add_heading("6.4 Sensitivity-Specificity Trade-off", level=2)
        self._add_image(plots.get("v15"), REPORT["img_md"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 7 -- Demographics
    # ══════════════════════════════════════════════════════════════════════

    def _section_7_demographics(self, demo_results, plots):
        self.doc.add_heading("7. Demographic Evaluation", level=1)

        self._interp(
            "This pipeline processes data from autistic children -- a population "
            "with significant inter-individual variability. A model that performs "
            "well on average but fails for severely autistic or non-verbal "
            "participants is clinically dangerous. This section breaks down "
            "performance by autism severity, verbal status, and age group."
        )

        self.doc.add_heading("7.1 Performance by Autism Severity", level=2)
        self._interp(
            "Autism severity (DSM-5 Levels 1-3) modulates physiological "
            "reactivity: more severe autism is associated with stronger EDA and "
            "HR responses (Kushki et al., 2013). Models may find it easier to "
            "classify emotions in severely autistic children due to larger "
            "physiological signals, potentially masking poor performance for "
            "Level 1 participants."
        )
        self._add_image(plots.get("v16"), REPORT["img_md"])

        self.doc.add_heading("7.2 Performance by Verbal Status", level=2)
        self._interp(
            "Non-verbal and minimally verbal participants are the primary "
            "clinical target -- they cannot self-report internal states. If the "
            "model performs worse for this subgroup, it fails the population "
            "it was designed to serve."
        )
        self._add_image(plots.get("v17"), REPORT["img_md"])

        if not demo_results:
            self._interp("No demographic data available for analysis.")

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 8 -- Equity
    # ══════════════════════════════════════════════════════════════════════

    def _section_8_equity(self, demo_results, plots):
        self.doc.add_heading("8. Equity Analysis", level=1)

        self._interp(
            "Equity analysis tests whether model performance varies significantly "
            "across demographic subgroups. Statistical tests (Mann-Whitney U for "
            "2 groups, Kruskal-Wallis for 3+) compare per-sample correctness "
            "distributions. Significant differences (p < 0.05) indicate the model "
            "may systematically disadvantage certain subgroups."
        )

        self.doc.add_heading("8.1 Equity Radar Chart", level=2)
        self._add_image(plots.get("v18"), REPORT["img_md"])

        self.doc.add_heading("8.2 Disparity Significance", level=2)
        self._add_image(plots.get("v19"), REPORT["img_md"])

        if demo_results:
            for name, dr in demo_results.items():
                if dr.flags:
                    self.doc.add_heading(f"Flags -- {name}", level=3)
                    for flag in dr.flags:
                        self._interp_bold_lead("WARNING: ", flag)

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 9 -- Robustness
    # ══════════════════════════════════════════════════════════════════════

    def _section_9_robustness(self, rob_results, plots):
        self.doc.add_heading("9. Robustness Analysis", level=1)

        self._interp(
            "Real-world wearable data is noisy: sensors shift, electrodes "
            "detach, motion artefacts corrupt signals. A robust model should "
            "degrade gracefully under perturbation. This section tests the "
            "top-performing models under three realistic failure modes: "
            "Gaussian noise injection, random feature dropout (missing data), "
            "and signal channel dropout (sensor failure)."
        )

        self.doc.add_heading("9.1 Noise Injection", level=2)
        self._interp(
            "Gaussian noise is added at increasing levels (1%-20% of each "
            "feature's standard deviation). The degradation curve shows how "
            "quickly performance drops. Steep decline indicates fragile models; "
            "gradual decline indicates robustness."
        )
        self._add_image(plots.get("v20"), REPORT["img_md"])

        self.doc.add_heading("9.2 Channel Dropout", level=2)
        self._interp(
            "Entire signal channels (EDA, BVP, ST, ACC, IBI) are zeroed out "
            "to simulate sensor failure. The F1 drop for each channel indicates "
            "the model's dependence on that signal. A model that collapses when "
            "EDA is lost may not be suitable for deployment where sensor "
            "attachment is unreliable."
        )
        self._add_image(plots.get("v21"), REPORT["img_md"])

        self.doc.add_heading("9.3 Missing Feature Tolerance", level=2)
        self._add_image(plots.get("v22"), REPORT["img_md"])

        self.doc.add_heading("9.4 Robustness Ranking", level=2)
        self._add_image(plots.get("v23"), REPORT["img_md"])

        if rob_results:
            for name, rr in rob_results.items():
                self._interp_bold_lead(
                    f"{name}: ",
                    f"Robustness score = {rr.robustness_score:.3f}. "
                    f"Worst channel dropout F1 loss = {rr.worst_channel_drop:.3f}."
                )

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 10 -- Statistical Comparison
    # ══════════════════════════════════════════════════════════════════════

    def _section_10_statistical(self, stat_comp, plots):
        self.doc.add_heading("10. Statistical Model Comparison", level=1)

        self._interp(
            "Raw performance differences between models may be due to chance. "
            "This section applies rigorous statistical tests to determine "
            "whether observed differences are statistically significant."
        )

        self.doc.add_heading("10.1 McNemar's Pairwise Tests", level=2)
        self._interp(
            "McNemar's test compares two classifiers by examining the 2x2 "
            "contingency table of their disagreements. It tests whether the "
            "two models make the same types of errors, regardless of their "
            "overall accuracy. A significant p-value indicates the models "
            "commit systematically different errors."
        )
        self._add_image(plots.get("v24"), REPORT["img_md"])

        if stat_comp and stat_comp.mcnemar_tests:
            sig_pairs = [t for t in stat_comp.mcnemar_tests if t.significant]
            self._interp(
                f"{len(sig_pairs)}/{len(stat_comp.mcnemar_tests)} pairs show "
                f"significantly different error patterns (p < 0.05)."
            )

        self.doc.add_heading("10.2 AUC-ROC Comparison", level=2)
        self._interp(
            "For models with probability outputs, AUC-ROC values are compared "
            "using bootstrap-based tests (DeLong's test for binary, bootstrap "
            "macro-AUC for multi-class). This tests whether one model's "
            "discriminative ability is genuinely superior."
        )
        self._add_image(plots.get("v25"), REPORT["img_md"])

        self.doc.add_heading("10.3 Cochran's Q Omnibus Test", level=2)
        self._interp(
            "Cochran's Q is a non-parametric omnibus test that simultaneously "
            "compares all models. A significant result means at least one model "
            "has a different error rate. It does not identify which model differs "
            "-- for that, consult the pairwise McNemar tests above."
        )
        self._add_image(plots.get("v26"), REPORT["img_sm"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 11 -- Error Analysis
    # ══════════════════════════════════════════════════════════════════════

    def _section_11_errors(self, error_results, plots):
        self.doc.add_heading("11. Error Analysis", level=1)

        self._interp(
            "Understanding why a model fails is as important as measuring how "
            "often it fails. This section identifies the most commonly confused "
            "class pairs, characterises failure modes, and examines whether "
            "errors cluster by participant."
        )

        self.doc.add_heading("11.1 Most Confused Class Pairs", level=2)
        self._add_image(plots.get("v27"), REPORT["img_md"])

        if error_results:
            best = list(error_results.values())[0]
            if best.confused_pairs:
                top = best.confused_pairs[0]
                self._interp_bold_lead(
                    f"Top confusion: ",
                    f"'{top.true_class}' -> '{top.pred_class}' accounts for "
                    f"{top.fraction_of_errors:.0%} of all errors "
                    f"({top.count} samples). "
                    f"This may reflect physiological similarity between these "
                    f"states (similar autonomic signatures) or insufficient "
                    f"training data for one of the classes."
                )

        self.doc.add_heading("11.2 Per-User Error Rates", level=2)
        self._interp(
            "If errors concentrate in certain participants, it suggests the "
            "model struggles with specific physiological profiles rather than "
            "failing uniformly. High per-user variance motivates person-specific "
            "model fine-tuning or transfer learning approaches."
        )
        self._add_image(plots.get("v28"), REPORT["img_md"])

        self.doc.add_heading("11.3 Misclassification Flow", level=2)
        self._add_image(plots.get("v29"), REPORT["img_md"])

        self.doc.add_heading("11.4 Failure Mode Visualisation", level=2)
        self._add_image(plots.get("v30"), REPORT["img_md"])

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 12 -- Inference Profiling
    # ══════════════════════════════════════════════════════════════════════

    def _section_12_inference(self, prof_results, plots):
        self.doc.add_heading("12. Inference Profiling", level=1)

        self._interp(
            "Clinical deployment requires models that can run within real-time "
            "constraints on the target hardware. For wearable-based monitoring, "
            "a single-sample prediction must complete within 100ms to keep pace "
            "with the 60-second feature extraction windows. This section profiles "
            "inference latency across different batch sizes."
        )

        self.doc.add_heading("12.1 Latency Comparison", level=2)
        self._add_image(plots.get("v31"), REPORT["img_md"])

        self.doc.add_heading("12.2 Pareto Frontier (F1 vs Latency)", level=2)
        self._interp(
            "The Pareto frontier identifies models that offer the best trade-off "
            "between prediction quality and speed. Models on the frontier cannot "
            "be improved in one dimension without sacrificing the other. The ideal "
            "deployment candidate sits in the upper-left quadrant: high F1, low "
            "latency."
        )
        self._add_image(plots.get("v32"), REPORT["img_md"])

        if prof_results:
            rt_models = [n for n, p in prof_results.items() if p.meets_realtime]
            self._interp_bold_lead(
                "Real-time capable: ",
                f"{len(rt_models)}/{len(prof_results)} models meet the <100ms "
                f"single-sample threshold: {', '.join(rt_models) if rt_models else 'none'}."
            )

        self.doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    #  Section 13 -- Summary & Recommendations
    # ══════════════════════════════════════════════════════════════════════

    def _section_13_summary(self, eval_results, gen_results, cal_results,
                            clinical_results, rob_results, prof_results):
        self.doc.add_heading("13. Summary & Recommendations", level=1)

        valid = [r for r in eval_results if not r.skipped]
        if not valid:
            self._interp("No models were successfully evaluated.")
            return

        self.doc.add_heading("13.1 Final Model Ranking", level=2)
        rows = []
        for i, r in enumerate(valid):
            row = {
                "Rank": i + 1,
                "Model": r.model_name,
                "Test F1": f"{r.test_metrics.get('f1_weighted', 0):.4f}",
            }
            if gen_results and r.model_name in gen_results:
                row["Generalisation"] = gen_results[r.model_name].verdict
            if cal_results and r.model_name in cal_results:
                row["Calibration"] = cal_results[r.model_name].verdict
            if prof_results and r.model_name in prof_results:
                row["Realtime"] = "YES" if prof_results[r.model_name].meets_realtime else "NO"
            rows.append(row)
        self._add_table(pd.DataFrame(rows))

        best = valid[0]
        self._interp_bold_lead(
            "Recommended model: ",
            f"{best.model_name} offers the best overall performance on the "
            f"held-out test set with F1 = {best.test_metrics.get('f1_weighted', 0):.4f}."
        )

        self.doc.add_heading("13.2 Clinical Deployment Readiness", level=2)
        self._interp(
            "Before clinical deployment, the following must be verified:\n"
            "1. Ethics board approval for real-time monitoring\n"
            "2. Participant/caregiver informed consent\n"
            "3. Model performance validated on real (non-simulated) data\n"
            "4. Alert thresholds calibrated with clinical input\n"
            "5. Failsafe mechanisms for sensor disconnection\n"
            "6. Data privacy and storage compliance (GDPR Article 9)"
        )

        self.doc.add_heading("13.3 Limitations", level=2)
        self._interp(
            "This evaluation was conducted on simulated physiological data. "
            "While the simulation models realistic autonomic responses based "
            "on published literature (Kreibig 2010, Kushki 2013, Stephens 2022), "
            "real-world data will introduce additional variability: motion "
            "artefacts during daily activities, sensor drift over extended "
            "wear periods, and medication effects on autonomic responses. "
            "All performance metrics reported here should be considered "
            "optimistic upper bounds until validated on clinical data."
        )

        self.doc.add_heading("13.4 Next Steps -> Module 9 (Deployment)", level=2)
        self._interp(
            "Module 9 will package the selected model for real-time inference:\n"
            "• Export to ONNX format for edge deployment\n"
            "• Build sliding-window inference engine\n"
            "• Implement caregiver alert system with confidence thresholds\n"
            "• Add graceful degradation for missing sensor channels\n"
            "• Deploy logging for retrospective analysis"
        )
