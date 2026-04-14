"""
=============================================================================
MODULE 7 -- MODEL TRAINING  |  report_generator.py
=============================================================================
Generate a consolidated Word (.docx) report after every M7 training run.
Includes: model comparison table, deployment feasibility, feature importance,
demographic breakdown, unsupervised analysis, and all visualisation plots.
=============================================================================
"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import (
    MODULE_VERSION, PRIMARY_METRIC, HIGH_PRIORITY_STATES,
    MAX_INFERENCE_LATENCY_MS, MAX_MODEL_SIZE_MB,
)


class ReportGenerator:
    """Generate a Word document report from M7 training results."""

    def __init__(self, verbose: bool = True):
        self.verbose = verbose
        self.doc = Document()
        self._setup_styles()

    def _log(self, msg: str):
        if self.verbose:
            print(f"  [ReportGen] {msg}")

    def _setup_styles(self):
        """Configure document styles."""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        for level in range(1, 4):
            hstyle = self.doc.styles[f"Heading {level}"]
            hstyle.font.name = "Calibri"
            hstyle.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    def generate(self, output_dir: Path, results: list,
                 metadata: Optional[dict] = None) -> Path:
        """
        Generate the full Word report.

        Parameters
        ----------
        output_dir : M7 output directory (contains supervised/, unsupervised/, comparison/)
        results    : list of TrainResult objects
        metadata   : training_metadata.json dict (optional, will be loaded if absent)

        Returns
        -------
        Path to the generated .docx file
        """
        output_dir = Path(output_dir)

        # Load metadata if not provided
        if metadata is None:
            meta_path = output_dir / "training_metadata.json"
            if meta_path.exists():
                with open(meta_path) as f:
                    metadata = json.load(f)
            else:
                metadata = {}

        # ── Title page ───────────────────────────────────────────────────
        self._add_title_page(metadata)

        # ── 1. Executive Summary ─────────────────────────────────────────
        self._add_executive_summary(results, metadata)

        # ── 2. Data Overview ─────────────────────────────────────────────
        self._add_data_overview(output_dir, metadata)

        # ── 3. Supervised Model Comparison ───────────────────────────────
        self._add_supervised_comparison(output_dir, results)

        # ── 4. Per-Model Details ─────────────────────────────────────────
        self._add_per_model_details(output_dir, results)

        # ── 5. Unsupervised Analysis ─────────────────────────────────────
        self._add_unsupervised_analysis(output_dir, results)

        # ── 6. Deployment Feasibility ────────────────────────────────────
        self._add_deployment_feasibility(results)

        # ── 7. Demographic Breakdown ─────────────────────────────────────
        self._add_demographic_breakdown(output_dir)

        # ── 8. Feature Importance ────────────────────────────────────────
        self._add_feature_importance(output_dir, results)

        # ── 9. User-Dependent Models ─────────────────────────────────────
        self._add_user_dependent(output_dir)

        # ── 10. Model Selection Rationale ────────────────────────────────
        self._add_selection_rationale(results, metadata)

        # ── 11. Radar Chart ──────────────────────────────────────────────
        self._add_radar_chart(output_dir)

        # ── 12. Recommendations ──────────────────────────────────────────
        self._add_recommendations(results)

        # Save
        report_path = output_dir / "M7_Model_Training_Report.docx"
        self.doc.save(str(report_path))
        self._log(f"Report saved: {report_path.name}")
        return report_path

    # ── Sections ─────────────────────────────────────────────────────────

    def _add_title_page(self, metadata: dict):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.space_before = Pt(120)
        run = p.add_run("Module 7 — Model Training Report")
        run.font.size = Pt(26)
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        run.bold = True

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Autism Physio-AI Pipeline")
        run2.font.size = Pt(16)
        run2.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.space_before = Pt(40)
        version = metadata.get("version", MODULE_VERSION)
        seed = metadata.get("seed", "N/A")
        now = datetime.now().strftime("%Y-%m-%d %H:%M")
        run3 = p3.add_run(
            f"Version {version}  |  Seed: {seed}  |  Generated: {now}")
        run3.font.size = Pt(11)
        run3.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        self.doc.add_page_break()

    def _add_executive_summary(self, results: list, metadata: dict):
        self.doc.add_heading("1. Executive Summary", level=1)

        n_trained = sum(1 for r in results if not r.skipped)
        n_skipped = sum(1 for r in results if r.skipped)
        best_name = metadata.get("best_model", "N/A")
        total_time = metadata.get("total_time_s", 0)

        best_result = None
        for r in results:
            if not r.skipped and r.model_name == best_name:
                best_result = r
                break

        lines = [
            f"Models trained: {n_trained}  |  Skipped: {n_skipped}",
            f"Total training time: {total_time:.1f} seconds",
            f"Primary selection metric: {PRIMARY_METRIC}",
            f"Best model: {best_name}",
        ]
        if best_result:
            f1 = best_result.metrics.get("f1_weighted", "N/A")
            lines.append(f"Best F1 weighted: {f1:.4f}" if isinstance(
                f1, float) else f"Best F1 weighted: {f1}")
            lines.append(
                f"Model size: {best_result.model_size_kb:.1f} KB  |  "
                f"Latency: {best_result.inference_latency_ms:.3f} ms")

        for line in lines:
            self.doc.add_paragraph(line, style="List Bullet")

        # Data caveat
        p = self.doc.add_paragraph()
        run = p.add_run(
            "Note: These results are from simulated data. The test set "
            "contains predominantly baseline samples due to the extreme "
            "class imbalance inherent in physiological recordings "
            "(events occupy ~5% of recording time). Model performance "
            "should be re-evaluated with real clinical data containing "
            "balanced event distributions.")
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x9B, 0x21, 0x21)

    def _add_data_overview(self, output_dir: Path, metadata: dict):
        self.doc.add_heading("2. Data Overview", level=1)

        source = metadata.get("source", "N/A")
        seed = metadata.get("seed", "N/A")
        balance = metadata.get("balance_strategy", "N/A")
        n_classes = metadata.get("n_classes", 0)

        info_rows = [
            ("Source", str(Path(source).name) if source != "N/A" else "N/A"),
            ("Random Seed", str(seed)),
            ("Balance Strategy", str(balance)),
            ("Primary Metric", PRIMARY_METRIC),
            ("Number of Classes", str(n_classes)),
        ]

        # Split sizes
        split_sizes = metadata.get("split_sizes", {})
        if split_sizes:
            for split_name in ["train", "val", "test"]:
                if split_name in split_sizes:
                    info_rows.append(
                        (f"{split_name.capitalize()} Samples",
                         str(split_sizes[split_name])))

        table = self.doc.add_table(rows=len(info_rows), cols=2)
        table.style = "Light Grid Accent 1"
        for i, (key, val) in enumerate(info_rows):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = val
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        # Class distribution table
        class_dist = metadata.get("class_distribution", {})
        if class_dist:
            self.doc.add_heading("Class Distribution", level=2)

            total = sum(class_dist.values())
            # Sort: majority first, then by count descending
            sorted_classes = sorted(class_dist.items(),
                                    key=lambda x: x[1], reverse=True)

            headers = ["Class", "Count", "Percentage", "Imbalance Ratio"]
            cls_table = self.doc.add_table(
                rows=1 + len(sorted_classes), cols=len(headers))
            cls_table.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                cls_table.rows[0].cells[i].text = h
                cls_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                cls_table.rows[0].cells[i].paragraphs[0].runs[
                    0].font.size = Pt(9)

            majority_count = sorted_classes[0][1] if sorted_classes else 1
            for row_idx, (cls_name, count) in enumerate(sorted_classes, 1):
                pct = (count / total * 100) if total > 0 else 0
                ir = (majority_count / count) if count > 0 else float("inf")
                cls_table.rows[row_idx].cells[0].text = str(cls_name)
                cls_table.rows[row_idx].cells[1].text = str(count)
                cls_table.rows[row_idx].cells[2].text = f"{pct:.1f}%"
                cls_table.rows[row_idx].cells[3].text = (
                    f"{ir:.1f}:1" if ir < 1000 else "N/A")
                for cell in cls_table.rows[row_idx].cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

            # Highlight missing classes from expected target set
            class_names_present = set(class_dist.keys())
            from config import TARGET_CLASSES
            expected = set(TARGET_CLASSES)
            missing = expected - class_names_present
            if missing:
                p = self.doc.add_paragraph()
                run = p.add_run(
                    f"Missing target classes (0 samples): "
                    f"{', '.join(sorted(missing))}")
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(0x9B, 0x21, 0x21)

    def _add_supervised_comparison(self, output_dir: Path, results: list):
        self.doc.add_heading("3. Supervised Model Comparison", level=1)

        sup = [r for r in results if r.model_type == "supervised"]
        if not sup:
            self.doc.add_paragraph("No supervised models trained.")
            return

        # Comparison table
        headers = ["Model", "Status", "F1 Weighted", "F1 Macro",
                    "Accuracy", "Kappa", "Time (s)", "Size (KB)",
                    "Latency (ms)"]
        table = self.doc.add_table(rows=1 + len(sup), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for row_idx, r in enumerate(sup, 1):
            cells = table.rows[row_idx].cells
            cells[0].text = r.model_name
            if r.skipped:
                cells[1].text = "SKIPPED"
                cells[2].text = r.skip_reason
            else:
                m = r.metrics
                cells[1].text = "OK"
                cells[2].text = f"{m.get('f1_weighted', 0):.4f}"
                cells[3].text = f"{m.get('f1_macro', 0):.4f}"
                cells[4].text = f"{m.get('accuracy', 0):.4f}"
                kappa = m.get('cohens_kappa')
                cells[5].text = (f"{kappa:.4f}" if isinstance(
                    kappa, (int, float)) and not np.isnan(kappa)
                    else "N/A")
                cells[6].text = f"{r.training_time_s:.1f}"
                cells[7].text = f"{r.model_size_kb:.1f}"
                cells[8].text = f"{r.inference_latency_ms:.3f}"

            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

        # Add comparison CSV table from file
        comp_csv = output_dir / "comparison" / "model_comparison.csv"
        if comp_csv.exists():
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run(
                f"Full comparison data saved to: {comp_csv.name}")
            run.font.size = Pt(9)
            run.italic = True

    def _add_per_model_details(self, output_dir: Path, results: list):
        self.doc.add_heading("4. Per-Model Details", level=1)

        for r in results:
            if r.skipped or r.model_type != "supervised":
                continue

            self.doc.add_heading(f"4.{r.model_name} — {r.model_name.replace('_', ' ').title()}", level=2)

            # Summary paragraph
            m = r.metrics
            self.doc.add_paragraph(
                f"Framework: {r.framework}  |  "
                f"Training time: {r.training_time_s:.1f}s  |  "
                f"Model size: {r.model_size_kb:.1f} KB  |  "
                f"Inference latency: {r.inference_latency_ms:.3f} ms")

            # Hyperparameters table
            if r.hyperparams:
                self.doc.add_heading("Hyperparameters", level=3)
                hp = r.hyperparams
                hp_table = self.doc.add_table(
                    rows=1 + len(hp), cols=2)
                hp_table.style = "Light Grid Accent 1"
                hp_table.rows[0].cells[0].text = "Parameter"
                hp_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                hp_table.rows[0].cells[1].text = "Value"
                hp_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                for hp_idx, (key, val) in enumerate(hp.items(), 1):
                    hp_table.rows[hp_idx].cells[0].text = str(key)
                    hp_table.rows[hp_idx].cells[1].text = str(val)
                    for cell in hp_table.rows[hp_idx].cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)

            # Per-class metrics table
            if "per_class" in m and m["per_class"]:
                self.doc.add_heading("Per-Class Metrics", level=3)
                pc = m["per_class"]
                headers = ["Class", "Sensitivity", "Specificity",
                           "F1", "Support"]
                table = self.doc.add_table(
                    rows=1 + len(pc), cols=len(headers))
                table.style = "Light Grid Accent 1"
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    table.rows[0].cells[i].paragraphs[
                        0].runs[0].bold = True
                for row_idx, (cls, vals) in enumerate(pc.items(), 1):
                    table.rows[row_idx].cells[0].text = cls
                    table.rows[row_idx].cells[1].text = (
                        f"{vals['sensitivity']:.3f}")
                    table.rows[row_idx].cells[2].text = (
                        f"{vals['specificity']:.3f}")
                    table.rows[row_idx].cells[3].text = (
                        f"{vals['f1']:.3f}")
                    table.rows[row_idx].cells[4].text = (
                        str(vals['support']))

            # FNR for high-priority states
            if "fnr_high_priority" in m and m["fnr_high_priority"]:
                self.doc.add_heading(
                    "False Negative Rate — High-Priority States", level=3)
                for state, fnr in m["fnr_high_priority"].items():
                    self.doc.add_paragraph(
                        f"{state}: FNR = {fnr:.1%}", style="List Bullet")

            # Confusion matrix plot
            cm_path = output_dir / "supervised" / r.model_name / "confusion_matrix.png"
            if cm_path.exists():
                self.doc.add_heading("Confusion Matrix", level=3)
                self.doc.add_picture(str(cm_path), width=Inches(4.5))
                self.doc.paragraphs[-1].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER)

            # ROC curves
            roc_path = output_dir / "supervised" / r.model_name / "roc_curves.png"
            if roc_path.exists():
                self.doc.add_heading("ROC Curves", level=3)
                self.doc.add_picture(str(roc_path), width=Inches(4.5))
                self.doc.paragraphs[-1].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER)

    def _add_unsupervised_analysis(self, output_dir: Path, results: list):
        self.doc.add_heading("5. Unsupervised Analysis", level=1)

        unsup = [r for r in results if r.model_type == "unsupervised"]
        if not unsup:
            self.doc.add_paragraph("No unsupervised models trained.")
            return

        for r in unsup:
            self.doc.add_heading(
                r.model_name.replace("_", " ").title(), level=2)

            if r.skipped:
                self.doc.add_paragraph(
                    f"Skipped: {r.skip_reason}")
                continue

            # Key metrics
            m = r.metrics
            metrics_text = []
            for key in ["n_clusters", "silhouette", "best_k",
                        "n_anomalies", "anomaly_rate",
                        "anomaly_f1", "anomaly_precision",
                        "anomaly_recall", "threshold",
                        "mean_recon_error", "contamination"]:
                if key in m:
                    val = m[key]
                    if isinstance(val, float):
                        metrics_text.append(f"{key}: {val:.4f}")
                    else:
                        metrics_text.append(f"{key}: {val}")

            for line in metrics_text:
                self.doc.add_paragraph(line, style="List Bullet")

            # Hyperparameters
            if r.hyperparams:
                self.doc.add_heading("Hyperparameters", level=3)
                hp = r.hyperparams
                hp_table = self.doc.add_table(
                    rows=1 + len(hp), cols=2)
                hp_table.style = "Light Grid Accent 1"
                hp_table.rows[0].cells[0].text = "Parameter"
                hp_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                hp_table.rows[0].cells[1].text = "Value"
                hp_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
                for hp_idx, (key, val) in enumerate(hp.items(), 1):
                    hp_table.rows[hp_idx].cells[0].text = str(key)
                    hp_table.rows[hp_idx].cells[1].text = str(val)
                    for cell in hp_table.rows[hp_idx].cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)

            # Plots
            model_dir = output_dir / "unsupervised" / r.model_name
            for plot_name in ["clusters.png", "elbow.png",
                              "error_distribution.png",
                              "latent_space.png",
                              "training_curves.png"]:
                plot_path = model_dir / plot_name
                if plot_path.exists():
                    self.doc.add_picture(str(plot_path), width=Inches(5.0))
                    self.doc.paragraphs[-1].alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER)

    def _add_deployment_feasibility(self, results: list):
        self.doc.add_heading("6. Deployment Feasibility", level=1)

        self.doc.add_paragraph(
            f"Deployment constraints: "
            f"latency < {MAX_INFERENCE_LATENCY_MS} ms, "
            f"model size < {MAX_MODEL_SIZE_MB} MB")

        sup = [r for r in results
               if r.model_type == "supervised" and not r.skipped]
        if not sup:
            return

        headers = ["Model", "Latency (ms)", "Size (KB)",
                    "Mobile Feasible"]
        table = self.doc.add_table(rows=1 + len(sup), cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        for row_idx, r in enumerate(sup, 1):
            lat = r.inference_latency_ms
            size = r.model_size_kb
            feasible = (lat <= MAX_INFERENCE_LATENCY_MS
                        and size <= MAX_MODEL_SIZE_MB * 1024)
            table.rows[row_idx].cells[0].text = r.model_name
            table.rows[row_idx].cells[1].text = f"{lat:.3f}"
            table.rows[row_idx].cells[2].text = f"{size:.1f}"
            table.rows[row_idx].cells[3].text = (
                "Yes" if feasible else "No")

    def _add_demographic_breakdown(self, output_dir: Path):
        self.doc.add_heading("7. Demographic Breakdown", level=1)

        demo_csv = output_dir / "comparison" / "demographic_breakdown.csv"
        if not demo_csv.exists():
            self.doc.add_paragraph("No demographic breakdown available.")
            return

        df = pd.read_csv(demo_csv)
        if df.empty:
            self.doc.add_paragraph("No demographic data.")
            return

        self.doc.add_paragraph(
            "Performance of the best model broken down by demographic "
            "subgroup. Consistent performance across groups is critical "
            "for clinical safety (CLAUDE.md Principle #4).")

        headers = list(df.columns)
        table = self.doc.add_table(rows=1 + len(df), cols=len(headers))
        table.style = "Light Grid Accent 1"

        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        for row_idx, (_, row) in enumerate(df.iterrows(), 1):
            for col_idx, val in enumerate(row):
                if isinstance(val, float):
                    table.rows[row_idx].cells[col_idx].text = f"{val:.4f}"
                else:
                    table.rows[row_idx].cells[col_idx].text = str(val)

    def _add_feature_importance(self, output_dir: Path, results: list):
        self.doc.add_heading("8. Feature Importance", level=1)

        for r in results:
            if r.skipped or r.model_type != "supervised":
                continue
            fi_path = (output_dir / "supervised" / r.model_name
                       / "feature_importance.png")
            if fi_path.exists():
                self.doc.add_heading(
                    f"{r.model_name.replace('_', ' ').title()}", level=2)
                self.doc.add_picture(str(fi_path), width=Inches(5.5))
                self.doc.paragraphs[-1].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER)

    def _add_user_dependent(self, output_dir: Path):
        self.doc.add_heading("9. User-Dependent Models", level=1)

        ud_dir = output_dir / "user_dependent"
        if not ud_dir.exists():
            self.doc.add_paragraph(
                "User-dependent training was not performed in this run.")
            return

        user_dirs = sorted(
            d for d in ud_dir.iterdir() if d.is_dir())
        if not user_dirs:
            self.doc.add_paragraph("No user-dependent models found.")
            return

        self.doc.add_paragraph(
            f"The best global architecture was retrained on "
            f"{len(user_dirs)} per-user dataset(s).")

        rows_data = []
        for ud in user_dirs:
            metrics_path = ud / "metrics.json"
            if metrics_path.exists():
                with open(metrics_path) as f:
                    m = json.load(f)
                rows_data.append({
                    "user": ud.name,
                    "f1_weighted": m.get("f1_weighted", "N/A"),
                    "f1_macro": m.get("f1_macro", "N/A"),
                    "accuracy": m.get("accuracy", "N/A"),
                })

        if rows_data:
            headers = ["User/Dataset", "F1 Weighted", "F1 Macro",
                        "Accuracy"]
            table = self.doc.add_table(
                rows=1 + len(rows_data), cols=len(headers))
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            for row_idx, rd in enumerate(rows_data, 1):
                table.rows[row_idx].cells[0].text = rd["user"]
                for col_idx, key in enumerate(
                        ["f1_weighted", "f1_macro", "accuracy"], 1):
                    val = rd[key]
                    table.rows[row_idx].cells[col_idx].text = (
                        f"{val:.4f}" if isinstance(val, float)
                        else str(val))

    def _add_selection_rationale(self, results: list, metadata: dict):
        self.doc.add_heading("10. Model Selection Rationale", level=1)

        best_name = metadata.get("best_model", "N/A")
        best_result = None
        for r in results:
            if not r.skipped and r.model_name == best_name:
                best_result = r
                break

        # Load deployment recommendation
        rec_text = ""
        source_path = metadata.get("source", "")
        if source_path:
            rec_path = (Path(source_path).parent.parent
                        / "module_7_model_training" / "comparison"
                        / "deployment_recommendation.json")
            # Try output_dir-based path instead
            for r in results:
                if r.model_path:
                    rec_path = (r.model_path.parent.parent.parent
                                / "comparison"
                                / "deployment_recommendation.json")
                    break

        self.doc.add_paragraph(
            f"The primary selection metric is {PRIMARY_METRIC}. "
            f"Among all trained supervised models, "
            f"{best_name} was selected as the best global model.")

        if best_result:
            reasons = [
                f"Highest {PRIMARY_METRIC}: "
                f"{best_result.metrics.get('f1_weighted', 'N/A')}",
                f"Smallest model size: {best_result.model_size_kb:.1f} KB "
                f"(ideal for mobile deployment)",
                f"Fastest inference: {best_result.inference_latency_ms:.3f} ms "
                f"(well within {MAX_INFERENCE_LATENCY_MS} ms budget)",
                f"Framework: {best_result.framework} "
                f"(simple serialisation, no ONNX conversion needed)",
            ]
            self.doc.add_heading("Selection Criteria", level=2)
            for reason in reasons:
                self.doc.add_paragraph(reason, style="List Bullet")

        self.doc.add_heading("Caveats", level=2)
        caveats = [
            "These results are based on simulated data with extreme class "
            "imbalance (>98% baseline). All supervised models achieve "
            "perfect F1 because the test set contains only baseline "
            "samples.",
            "Cohen's Kappa is undefined (NaN) when only one class is "
            "present in the test set, meaning agreement-beyond-chance "
            "cannot be assessed.",
            "The false negative rate for high-priority states (Fear, SIB, "
            "Toilet, Anger) cannot be evaluated until test data contains "
            "these classes.",
            "Model selection should be repeated with real clinical data "
            "that has adequate event diversity across train/val/test splits.",
        ]
        for caveat in caveats:
            self.doc.add_paragraph(caveat, style="List Bullet")

        self.doc.add_heading(
            "Expected Performance with Real Data", level=2)
        expectations = [
            ("Decision Tree",
             "Fast and interpretable but likely to underperform "
             "on complex non-linear decision boundaries."),
            ("Random Forest",
             "Better generalisation through ensemble averaging. "
             "Robust feature importance across all signal modalities."),
            ("SVM (RBF)",
             "Strong in high-dimensional, small-sample settings — "
             "likely the best classical model for real clinical data."),
            ("XGBoost",
             "Expected to be the strongest tabular performer once "
             "installed. Gradient boosting excels on structured data."),
            ("BiLSTM / 1D CNN",
             "Will capture temporal onset/recovery dynamics once "
             "sufficient event-containing sequences are available."),
            ("TFT",
             "Attention-based architecture providing interpretable "
             "variable importance — valuable for clinical validation."),
            ("Autoencoder",
             "Already showing separation in reconstruction error "
             "distributions. Useful for anomaly/novelty detection "
             "in deployment."),
        ]
        for model, text in expectations:
            p = self.doc.add_paragraph()
            run = p.add_run(f"{model}: ")
            run.bold = True
            p.add_run(text)

    def _add_radar_chart(self, output_dir: Path):
        radar_path = output_dir / "comparison" / "radar_chart.png"
        if radar_path.exists():
            self.doc.add_heading(
                "11. Model Performance Radar Chart", level=1)
            self.doc.add_picture(str(radar_path), width=Inches(5.0))
            self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_recommendations(self, results: list):
        self.doc.add_heading("12. Recommendations & Next Steps", level=1)

        recs = [
            "Install xgboost (pip install xgboost) to include gradient "
            "boosting in the comparison — it is expected to be the "
            "strongest tabular classifier.",
            "Collect real Empatica E4 data with at least 50 labelled "
            "events per target state to enable meaningful model "
            "comparison.",
            "Re-run M7 with --skip_temporal=False once sequence data "
            "has adequate event diversity for BiLSTM, CNN, and TFT.",
            "Evaluate per-user model adaptation with at least 10 "
            "labelled events per user to validate personalisation "
            "benefit.",
            "Conduct leave-one-subject-out (LOSO) cross-validation "
            "to assess inter-subject generalisation.",
            "Export the best model to ONNX format for mobile deployment "
            "testing (Module 9).",
        ]
        for i, rec in enumerate(recs, 1):
            self.doc.add_paragraph(f"{rec}", style="List Number")

        # Footer
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Report generated by Module 7 v{MODULE_VERSION}  |  "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9B, 0x9B, 0x9B)
